#check

import sys
import sqlite3
from PyQt5.QtWidgets import *
from PyQt5.uic import loadUi
from PyQt5.QtGui import QIcon, QPixmap
from docx import Document
import base64


class PlanDialog(QDialog):
    def __init__(self, serial_id, parent=None):
        super().__init__(parent)
        self.serial_id = serial_id
        self.initUI()

    def initUI(self):
        self.setWindowTitle('Запланировать просмотр')
        layout = QVBoxLayout()
        self.calendar_widget = QCalendarWidget()
        self.plan_button = QPushButton("Запланировать")
        self.plan_button.clicked.connect(self.plan_serial)
        layout.addWidget(self.calendar_widget)
        layout.addWidget(self.plan_button)
        self.setLayout(layout)

    def plan_serial(self):
        selected_date = self.calendar_widget.selectedDate()
        if selected_date.isValid():
            planned_date = selected_date.toString("yyyy-MM-dd")
            self.parent().add_to_table('planned', self.serial_id, planned_date)
            self.accept()


class MainWindow(QMainWindow):

    def __init__(self):
        super().__init__()
        loadUi('test.ui', self)
        self.con = sqlite3.connect('serial.sqlite')

        #settings
        self.encoded_image = None
        self.pushButton.clicked.connect(self.add_serial)
        self.pushButton_3.clicked.connect(self.delete_serial)
        self.pushButton_4.clicked.connect(self.update_serial)
        self.pushButton_2.clicked.connect(self.browse_files)
        self.pushButton_5.clicked.connect(self.filec)

        #list of serials
        self.show_serials()
        self.page_2 = self.stackedWidget.findChild(QWidget, 'page_2')
        self.layout_page_2 = QVBoxLayout(self.page_2)
        self.scroll_area = QScrollArea(self.page_2)
        self.scroll_area.setWidgetResizable(True)
        self.scroll_widget = QWidget()
        self.scroll_layout = QVBoxLayout(self.scroll_widget)
        self.scroll_area.setWidget(self.scroll_widget)
        self.layout_page_2.addWidget(self.scroll_area)
        self.add_checkboxes_to_page_2()

        #watching
        self.page_3 = self.stackedWidget.findChild(QWidget, 'page_3')
        self.layout_page_3 = QVBoxLayout(self.page_3)
        self.scroll_area_3 = QScrollArea(self.page_3)
        self.scroll_area_3.setWidgetResizable(True)
        self.scroll_widget_3 = QWidget()
        self.scroll_area_3.setWidget(self.scroll_widget_3)
        self.layout_page_3.addWidget(self.scroll_area_3)
        self.scroll_layout_3 = QVBoxLayout(self.scroll_widget_3)
        self.add_checkboxes_to_page_3()

        #watched
        self.page_5 = self.stackedWidget.findChild(QWidget, 'page_5')
        self.layout_page_5 = QVBoxLayout(self.page_5)
        self.scroll_area_5 = QScrollArea(self.page_5)
        self.scroll_area_5.setWidgetResizable(True)
        self.scroll_widget_5 = QWidget()
        self.scroll_area_5.setWidget(self.scroll_widget_5)
        self.layout_page_5.addWidget(self.scroll_area_5)
        self.scroll_layout_5 = QVBoxLayout(self.scroll_widget_5)
        self.add_checkboxes_to_page_5()

        #planned
        self.page_4 = self.stackedWidget.findChild(QWidget, 'page_4')
        self.layout_page_4 = QVBoxLayout(self.page_4)
        self.scroll_area_4 = QScrollArea(self.page_4)
        self.scroll_area_4.setWidgetResizable(True)
        self.scroll_widget_4 = QWidget()
        self.scroll_area_4.setWidget(self.scroll_widget_4)
        self.layout_page_4.addWidget(self.scroll_area_4)
        self.scroll_layout_4 = QVBoxLayout(self.scroll_widget_4)
        self.add_checkboxes_to_page_4()


    def add_checkboxes_to_page_2(self):
        for i in reversed(range(self.scroll_layout.count())):
            widget = self.scroll_layout.itemAt(i).widget()
            if widget is not None:
                widget.deleteLater()

        try:
            cursor = self.con.cursor()
            cursor.execute("SELECT id, name, genre, country, year, seasons, episodes, pic FROM serials")
            serials = cursor.fetchall()

            for serial_id, name, genre, country, year, seasons, episodes, pic in serials:
                serial_widget = QWidget()
                serial_layout = QHBoxLayout(serial_widget)

                image_label = QLabel()
                if pic is not None:
                    image_data = base64.b64decode(pic)
                else:
                    with open('zat3.jpg', 'rb') as f:
                        image_data = f.read()
                pixmap = QPixmap()
                pixmap.loadFromData(image_data)
                image_label.setPixmap(pixmap)
                serial_layout.addWidget(image_label)
                result = cursor.execute("SELECT name FROM genres WHERE id = ?",
                                        (genre,)).fetchall()
                info_layout = QVBoxLayout()

                info_label = QLabel(f"Название: {name}\n"
                                    f"Жанр: {result[0][0]}\n"
                                    f"Страна: {country}\n"
                                    f"Год: {year}\n"
                                    f"Сезоны: {seasons}\n"
                                    f"Эпизоды: {episodes}")
                info_layout.addWidget(info_label)

                checkbox_layout = QVBoxLayout()

                checkbox_planned = QCheckBox("Планирую посмотреть")
                checkbox_watched = QCheckBox("Просмотрено")
                checkbox_watching = QCheckBox("Смотрю")

                checkbox_planned.serial_id = serial_id
                checkbox_watched.serial_id = serial_id
                checkbox_watching.serial_id = serial_id

                if self.is_serial_in_planned(serial_id):
                    checkbox_planned.setChecked(True)
                    checkbox_watched.setEnabled(False)
                    checkbox_watching.setEnabled(False)
                else:
                    checkbox_watched.setEnabled(True)
                    checkbox_watching.setEnabled(True)
                    if self.is_serial_in_watched(serial_id):
                        checkbox_watched.setChecked(True)
                        checkbox_planned.setEnabled(False)
                        checkbox_watching.setEnabled(False)
                    else:
                        checkbox_planned.setEnabled(True)
                        checkbox_watching.setEnabled(True)
                        if self.is_serial_in_watching(serial_id):
                            checkbox_watching.setChecked(True)
                            checkbox_watched.setEnabled(False)
                            checkbox_planned.setEnabled(False)
                        else:
                            checkbox_watched.setEnabled(True)
                            checkbox_planned.setEnabled(True)

                checkbox_planned.stateChanged.connect(self.checkbox_changed)
                checkbox_watched.stateChanged.connect(self.checkbox_changed)
                checkbox_watching.stateChanged.connect(self.checkbox_changed)

                checkbox_layout.addWidget(checkbox_planned)
                checkbox_layout.addWidget(checkbox_watched)
                checkbox_layout.addWidget(checkbox_watching)

                info_layout.addLayout(checkbox_layout)
                serial_layout.addLayout(info_layout)
                serial_widget.setLayout(serial_layout)

                self.scroll_layout.addWidget(serial_widget)


        except sqlite3.Error as e:
            print("Произошла ошибка при получении данных из базы данных:", e)

        finally:
            cursor.close()

    def checkbox_changed(self, state):
        checkbox = self.sender()
        serial_id = checkbox.serial_id

        if checkbox.isChecked():
            if checkbox.text() == "Планирую посмотреть":
                dialog = PlanDialog(serial_id, parent=self)
                dialog.exec_()
            elif checkbox.text() == "Просмотрено":
                self.add_to_table('watched', serial_id)
            elif checkbox.text() == "Смотрю":
                self.add_to_table('watching', serial_id)
        else:
            if checkbox.text() == "Планирую посмотреть":
                self.remove_from_table('planned', serial_id)
            elif checkbox.text() == "Просмотрено":
                self.remove_from_table('watched', serial_id)
            elif checkbox.text() == "Смотрю":
                self.remove_from_table('watching', serial_id)

        self.add_checkboxes_to_page_2()
        self.add_checkboxes_to_page_3()
        self.add_checkboxes_to_page_4()
        self.add_checkboxes_to_page_5()

    def is_serial_in_planned(self, serial_id):
        try:
            cursor = self.con.cursor()
            cursor.execute("SELECT idSerial FROM planned WHERE idSerial = ?", (serial_id,))
            result = cursor.fetchone()
            return result is not None
        except sqlite3.Error as e:
            print("Произошла ошибка при проверке planned:", e)
        finally:
            cursor.close()

    def is_serial_in_watched(self, serial_id):
        try:
            cursor = self.con.cursor()
            cursor.execute("SELECT idSerial FROM watched WHERE idSerial = ?", (serial_id,))
            result = cursor.fetchone()
            return result is not None
        except sqlite3.Error as e:
            print("Произошла ошибка при проверке planned:", e)
        finally:
            cursor.close()

    def is_serial_in_watching(self, serial_id):
        try:
            cursor = self.con.cursor()
            cursor.execute("SELECT idSerial FROM watching WHERE idSerial = ?", (serial_id,))
            result = cursor.fetchone()
            return result is not None
        except sqlite3.Error as e:
            print("Произошла ошибка при проверке planned:", e)
        finally:
            cursor.close()

    def add_to_table(self, table_name, serial_id, date=None):
        try:
            cursor = self.con.cursor()
            print(table_name)
            if table_name == "planned":
                cursor.execute(f"INSERT INTO {table_name} (idSerial, date) VALUES (?, ?)", (idSerial := serial_id, date := date))
            else:
                cursor.execute(f"INSERT INTO {table_name} (idSerial) VALUES (?)", (idSerial := serial_id,))
            self.con.commit()
        except sqlite3.Error as e:
            print(f"Произошла ошибка при добавлении в таблицу {table_name}: {e}")
        finally:
            cursor.close()

    def remove_from_table(self, table_name, serial_id):
        try:
            cursor = self.con.cursor()
            cursor.execute(f"DELETE FROM {table_name} WHERE idSerial = ?", (serial_id,))
            self.con.commit()
        except sqlite3.Error as e:
            print(f"Произошла ошибка при удалении из таблицы {table_name}: {e}")
        finally:
            cursor.close()

    def add_checkboxes_to_page_3(self):  #watching
        for i in reversed(range(self.scroll_layout_3.count())):
            widget = self.scroll_layout_3.itemAt(i).widget()
            if widget is not None:
                widget.deleteLater()
        self.scroll_layout_3.addWidget(self.pushButton_5)
        try:
            cursor = self.con.cursor()
            cursor.execute("SELECT serials.id, serials.name, serials.genre, serials.country, serials.year, serials.seasons, serials.episodes, serials.pic FROM serials INNER JOIN watching ON serials.id = watching.idSerial")
            serials = cursor.fetchall()

            for serial_id, name, genre, country, year, seasons, episodes, pic in serials:
                serial_widget = QWidget()
                serial_layout = QHBoxLayout(serial_widget)

                image_label = QLabel()
                if pic is not None:
                    image_data = base64.b64decode(pic)
                else:
                    with open('zat3.jpg', 'rb') as f:
                        image_data = f.read()
                pixmap = QPixmap()
                pixmap.loadFromData(image_data)
                image_label.setPixmap(pixmap)
                serial_layout.addWidget(image_label)
                result = cursor.execute("SELECT name FROM genres WHERE id = ?",
                                        (genre,)).fetchall()

                info_layout = QVBoxLayout()

                info_label = QLabel(f"Название: {name}\n"
                                    f"Жанр: {result[0][0]}\n"
                                    f"Страна: {country}\n"
                                    f"Год: {year}\n"
                                    f"Сезоны: {seasons}\n"
                                    f"Эпизоды: {episodes}")
                info_layout.addWidget(info_label)

                checkbox_layout = QVBoxLayout()

                checkbox_planned = QCheckBox("Планирую посмотреть")
                checkbox_watched = QCheckBox("Просмотрено")
                checkbox_watching = QCheckBox("Смотрю")

                checkbox_planned.serial_id = serial_id
                checkbox_watched.serial_id = serial_id
                checkbox_watching.serial_id = serial_id


                if self.is_serial_in_planned(serial_id):
                    checkbox_planned.setChecked(True)
                    checkbox_watched.setCheckable(False)
                    checkbox_watching.setCheckable(False)
                if self.is_serial_in_watched(serial_id):
                    checkbox_watched.setChecked(True)
                    checkbox_planned.setCheckable(False)
                    checkbox_watching.setCheckable(False)
                if self.is_serial_in_watching(serial_id):
                    checkbox_watching.setChecked(True)
                    checkbox_watched.setCheckable(False)
                    checkbox_planned.setCheckable(False)

                checkbox_planned.stateChanged.connect(self.checkbox_changed)
                checkbox_watched.stateChanged.connect(self.checkbox_changed)
                checkbox_watching.stateChanged.connect(self.checkbox_changed)

                checkbox_layout.addWidget(checkbox_planned)
                checkbox_layout.addWidget(checkbox_watched)
                checkbox_layout.addWidget(checkbox_watching)

                info_layout.addLayout(checkbox_layout)
                serial_layout.addLayout(info_layout)
                serial_widget.setLayout(serial_layout)

                self.scroll_layout_3.addWidget(serial_widget)


        except sqlite3.Error as e:
            print("Произошла ошибка при получении данных из базы данных:", e)

        finally:
            cursor.close()

    def add_checkboxes_to_page_5(self): #watched
        for i in reversed(range(self.scroll_layout_5.count())):
            widget = self.scroll_layout_5.itemAt(i).widget()
            if widget is not None:
                widget.deleteLater()
        self.scroll_layout_5.addWidget(self.pushButton_5)
        try:
            cursor = self.con.cursor()
            cursor.execute("SELECT serials.id, serials.name, serials.genre, serials.country, serials.year, serials.seasons, serials.episodes, serials.pic FROM serials INNER JOIN watched ON serials.id = watched.idSerial")
            serials = cursor.fetchall()

            for serial_id, name, genre, country, year, seasons, episodes, pic in serials:
                serial_widget = QWidget()
                serial_layout = QHBoxLayout(serial_widget)

                image_label = QLabel()
                if pic is not None:
                    image_data = base64.b64decode(pic)
                else:
                    with open('zat3.jpg', 'rb') as f:
                        image_data = f.read()
                pixmap = QPixmap()
                pixmap.loadFromData(image_data)
                image_label.setPixmap(pixmap)
                serial_layout.addWidget(image_label)
                result = cursor.execute("SELECT name FROM genres WHERE id = ?",
                                        (genre,)).fetchall()

                info_layout = QVBoxLayout()

                info_label = QLabel(f"Название: {name}\n"
                                    f"Жанр: {result[0][0]}\n"
                                    f"Страна: {country}\n"
                                    f"Год: {year}\n"
                                    f"Сезоны: {seasons}\n"
                                    f"Эпизоды: {episodes}")
                info_layout.addWidget(info_label)

                checkbox_layout = QVBoxLayout()

                checkbox_planned = QCheckBox("Планирую посмотреть")
                checkbox_watched = QCheckBox("Просмотрено")
                checkbox_watching = QCheckBox("Смотрю")

                checkbox_planned.serial_id = serial_id
                checkbox_watched.serial_id = serial_id
                checkbox_watching.serial_id = serial_id

                if self.is_serial_in_planned(serial_id):
                    checkbox_planned.setChecked(True)
                if self.is_serial_in_watched(serial_id):
                    checkbox_watched.setChecked(True)
                if self.is_serial_in_watching(serial_id):
                    checkbox_watching.setChecked(True)

                checkbox_planned.stateChanged.connect(self.checkbox_changed)
                checkbox_watched.stateChanged.connect(self.checkbox_changed)
                checkbox_watching.stateChanged.connect(self.checkbox_changed)

                checkbox_layout.addWidget(checkbox_planned)
                checkbox_layout.addWidget(checkbox_watched)
                checkbox_layout.addWidget(checkbox_watching)

                info_layout.addLayout(checkbox_layout)
                serial_layout.addLayout(info_layout)
                serial_widget.setLayout(serial_layout)

                self.scroll_layout_5.addWidget(serial_widget)

        except sqlite3.Error as e:
            print("Произошла ошибка при получении данных из базы данных:", e)

        finally:
            cursor.close()

    def add_checkboxes_to_page_4(self): #planned
        for i in reversed(range(self.scroll_layout_4.count())):
            widget = self.scroll_layout_4.itemAt(i).widget()
            if widget is not None:
                widget.deleteLater()
        self.scroll_layout_4.addWidget(self.pushButton_5)
        try:
            cursor = self.con.cursor()
            cursor.execute("SELECT serials.id, serials.name, serials.genre, serials.country, serials.year, serials.seasons, serials.episodes, serials.pic, planned.date FROM serials INNER JOIN planned ON serials.id = planned.idSerial")
            serials = cursor.fetchall()

            for serial_id, name, genre, country, year, seasons, episodes, pic, date in serials:
                serial_widget = QWidget()
                serial_layout = QHBoxLayout(serial_widget)

                image_label = QLabel()
                if pic is not None:
                    image_data = base64.b64decode(pic)
                else:
                    with open('zat3.jpg', 'rb') as f:
                        image_data = f.read()
                pixmap = QPixmap()
                pixmap.loadFromData(image_data)
                image_label.setPixmap(pixmap)
                serial_layout.addWidget(image_label)
                result = cursor.execute("SELECT name FROM genres WHERE id = ?",
                                        (genre,)).fetchall()

                info_layout = QVBoxLayout()

                info_label = QLabel(f"Название: {name}\n"
                                    f"Жанр: {result[0][0]}\n"
                                    f"Страна: {country}\n"
                                    f"Год: {year}\n"
                                    f"Сезоны: {seasons}\n"
                                    f"Эпизоды: {episodes}\n"
                                    f"План: {date}")
                info_layout.addWidget(info_label)

                checkbox_layout = QVBoxLayout()

                start_watching = QPushButton("Начать просмотр")
                start_watching.clicked.connect(lambda: self.start(serial_id))
                checkbox_layout.addWidget(start_watching)
                info_layout.addLayout(checkbox_layout)
                serial_layout.addLayout(info_layout)
                serial_widget.setLayout(serial_layout)

                self.scroll_layout_4.addWidget(serial_widget)

        except sqlite3.Error as e:
            print("Произошла ошибка при получении данных из базы данных:", e)

        finally:
            cursor.close()

    def start(self, serialid):
        self.remove_from_table('planned', serialid)
        self.add_to_table('watching', serialid)
        self.add_checkboxes_to_page_3()
        self.add_checkboxes_to_page_4()
        self.add_checkboxes_to_page_2()



    def filec(self):
        document = Document()
        document.add_heading('Запланированно')
        cursor = self.con.cursor()
        cursor.execute(
            "SELECT serials.id, serials.name, serials.genre, serials.country, serials.year, serials.seasons, serials.episodes, serials.pic, planned.date FROM serials INNER JOIN planned ON serials.id = planned.idSerial")
        serials = cursor.fetchall()

        for serial_id, name, genre, country, year, seasons, episodes, pic, date in serials:
            result = cursor.execute("SELECT name FROM genres WHERE id = ?",
                                    (genre,)).fetchall()
            text_to_save = f"{name}: {result[0][0]}, {country}, {year}, {seasons}, {episodes}, {date} \n"
            document.add_paragraph(text_to_save)
        document.save('output.docx')


    #settings
    def browse_files(self):
        file_dialog = QFileDialog()
        file_dialog.setNameFilter("Images (*.jpg *.png *.webp *.gif)")
        file_dialog.setViewMode(QFileDialog.List)

        if file_dialog.exec_():
            selected_files = file_dialog.selectedFiles()
            if selected_files:
                file_path = selected_files[0]

            with open(file_path, 'rb') as file:
                image_data = file.read()

        # Кодируем изображение в base64
            self.encoded_image = base64.b64encode(image_data)

    def add_serial(self):
        try:
            cursor = self.con.cursor()
            result = cursor.execute("SELECT id FROM genres WHERE name = ?", (self.comboBox.currentText(),)).fetchall()
            cursor.execute("""INSERT INTO
                                        serials (name, genre, episodes, seasons, year, country, score, pic)
                                        VALUES(?, ?, ?, ?, ?, ?, ?, ?)""",
                        (name := self.lineEdit.text(), genre := result[0][0], episodes := self.spinBox.text(),
                         seasons := self.spinBox_2.text(), year := self.lineEdit_5.text(), country := self.lineEdit_6.text(),
                         score := self.horizontalSlider.value(), pic := self.encoded_image)).fetchall()
            self.con.commit()
            print("Данные успешно добавлены в базу данных!")
            QMessageBox.information(self, "Успех", "Данные успешно добавлены в базу данных!")
            self.show_serials()
            self.add_checkboxes_to_page_2()
        except sqlite3.Error as e:
            QMessageBox.critical(self, "Ошибка", f"Произошла ошибка при добавлении данных в базу данных: {e}")
        finally:
            cursor.close()

    def delete_serial(self):
        try:
            cursor = self.con.cursor()
            name = self.lineEdit_3.text()
            cursor.execute("DELETE FROM serials WHERE name = ?", (name,))
            self.con.commit()
            print(f"Запись с названием '{name}' успешно удалена из базы данных!")
            QMessageBox.information(self, "Успех", f"Запись с названием '{name}' успешно удалена из базы данных!")
            self.show_serials()
            self.add_checkboxes_to_page_2()
        except sqlite3.Error as e:
            QMessageBox.critical(self, "Ошибка", f"Произошла ошибка при удалении данных из базы данных: {e}")
        finally:
            cursor.close()

    def update_serial(self):
        try:
            cursor = self.con.cursor()
            new_genre = cursor.execute("SELECT id FROM genres WHERE name = ?", (self.comboBox.currentText(),)).fetchall()
            new_episodes = self.spinBox.value()
            new_seasons = self.spinBox_2.value()
            new_year = self.lineEdit_5.text()
            new_country = self.lineEdit_6.text()
            new_score = self.horizontalSlider.value()
            new_pic = self.encoded_image

            name = self.lineEdit.text()

            cursor.execute("""UPDATE serials
                              SET genre = ?, episodes = ?, seasons = ?, year = ?, country = ?, score = ?, pic = ?
                              WHERE name = ?""",
                           (new_genre[0][0], new_episodes, new_seasons, new_year, new_country, new_score, new_pic, name))
            self.con.commit()

            print(f"Запись с названием '{name}' успешно обновлена в базе данных!")
            QMessageBox.information(self, "Успех", f"Запись с названием '{name}' успешно обновлена в базе данных!")
            self.show_serials()
            self.add_checkboxes_to_page_2()
        except sqlite3.Error as e:
            QMessageBox.critical(self, "Ошибка", f"Произошла ошибка при обновлении данных в базе данных: {e}")

        finally:
            cursor.close()

    def show_serials(self):
        try:
            cursor = self.con.cursor()
            cursor.execute("SELECT name, genre, episodes, seasons, year, country, score FROM serials")
            serials = cursor.fetchall()

            self.tableWidget_2.setRowCount(len(serials))
            self.tableWidget_2.setColumnCount(7)

            self.tableWidget_2.setHorizontalHeaderLabels(["Название", "Жанр", "Эпизоды", "Сезоны", "Год", "Страна", "Рейтинг"])

            for row, serial in enumerate(serials):
                for col, data in enumerate(serial):
                    item = QTableWidgetItem(str(data))
                    self.tableWidget_2.setItem(row, col, item)

        except sqlite3.Error as e:
            print("Произошла ошибка при получении данных из базы данных:", e)

        finally:
            cursor.close()

if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec_())